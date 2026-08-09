"""Word (.docx) loader.

Extracts paragraphs and tables. Tables are flattened to pipe-separated
rows rather than dropped — spec sheets and comparison tables are often
the highest-value content in a Word document, and losing them silently
would make retrieval look broken.

Requires ``python-docx``. Legacy ``.doc`` (OLE2) is not supported; that
format needs an external converter and is out of scope.
"""

from __future__ import annotations

import asyncio
import zipfile
from pathlib import Path
from typing import Any

from ..models import Document
from .base import DocumentLoader, LoaderError


class DocxLoader(DocumentLoader):
    """Loads ``.docx`` files into flat text with headings preserved."""

    extensions = (".docx",)
    name = "docx"

    def _extract_sync(self, path: Path) -> tuple[str, dict[str, Any]]:
        try:
            import docx  # noqa: PLC0415
        except ImportError as exc:  # pragma: no cover
            raise LoaderError(
                "DOCX support requires python-docx. Install it with: uv add python-docx"
            ) from exc

        try:
            document = docx.Document(str(path))
        except zipfile.BadZipFile as exc:
            raise LoaderError(
                f"{path.name} is not a valid .docx file. Legacy .doc files must be "
                "converted to .docx before they can be indexed."
            ) from exc
        except Exception as exc:  # noqa: BLE001
            raise LoaderError(f"Could not open {path.name}: {exc}") from exc

        blocks: list[str] = []
        heading_count = 0
        for paragraph in document.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue
            style = getattr(paragraph.style, "name", "") or ""
            if style.startswith("Heading"):
                # Re-emit as Markdown so the chunker's paragraph splitting
                # and the LLM both see the document's structure.
                level = "".join(ch for ch in style if ch.isdigit()) or "1"
                blocks.append(f"{'#' * min(int(level), 6)} {text}")
                heading_count += 1
            else:
                blocks.append(text)

        table_count = 0
        for table in document.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [(cell.text or "").strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                table_count += 1
                blocks.append("\n".join(rows))

        info: dict[str, Any] = {
            "paragraph_count": len(document.paragraphs),
            "heading_count": heading_count,
            "table_count": table_count,
        }
        try:
            core = document.core_properties
            if core.title:
                info["docx_title"] = str(core.title)
            if core.author:
                info["docx_author"] = str(core.author)
        except Exception:  # noqa: BLE001
            pass

        return "\n\n".join(blocks), info

    async def load(
        self,
        path: Path,
        *,
        title: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> Document:
        content, info = await asyncio.to_thread(self._extract_sync, path)

        if not content.strip():
            raise LoaderError(f"{path.name} contains no readable text")

        return self.build_document(
            path=path,
            content=content,
            title=title or info.get("docx_title") or path.stem,
            metadata=metadata,
            extra=info,
        )
